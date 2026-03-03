"""
Export utilities: TXT, SRT, VTT, DOCX, JSON, CSV
"""

import json
import csv
import io
from datetime import datetime


def seconds_to_srt_time(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds - int(seconds)) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def seconds_to_vtt_time(seconds: float) -> str:
    return seconds_to_srt_time(seconds).replace(",", ".")


def export_as_txt(
    transcript: str,
    title: str = "",
    summary_text: str = "",
    include_summary: bool = False,
) -> bytes:
    header = f"# {title}\n# Exported: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n" if title else ""
    body = transcript or ""
    if include_summary and summary_text:
        body = f"{body}\n\n---\n\nSummary\n\n{summary_text}"
    return (header + body).encode("utf-8")


def export_as_json(transcription_data: dict, include_summary: bool = False) -> bytes:
    output = {
        "title": transcription_data.get("original_filename", ""),
        "model": transcription_data.get("model_used", ""),
        "language": transcription_data.get("language", ""),
        "duration_seconds": transcription_data.get("duration_seconds", 0),
        "word_count": transcription_data.get("word_count", 0),
        "created_at": transcription_data.get("created_at", ""),
        "transcript": transcription_data.get("transcript", ""),
    }
    if include_summary:
        output["summary"] = transcription_data.get("summary_text") or transcription_data.get("summary", "")
    return json.dumps(output, ensure_ascii=False, indent=2).encode("utf-8")


def export_as_srt(segments: list) -> bytes:
    """Export segments as SRT subtitles."""
    lines = []
    for i, seg in enumerate(segments, 1):
        start = seconds_to_srt_time(seg.get("start", 0))
        end = seconds_to_srt_time(seg.get("end", 0))
        text = seg.get("text", "").strip()
        lines.append(f"{i}\n{start} --> {end}\n{text}\n")
    return "\n".join(lines).encode("utf-8")


def export_as_vtt(segments: list) -> bytes:
    """Export segments as WebVTT."""
    lines = ["WEBVTT\n"]
    for seg in segments:
        start = seconds_to_vtt_time(seg.get("start", 0))
        end = seconds_to_vtt_time(seg.get("end", 0))
        text = seg.get("text", "").strip()
        lines.append(f"{start} --> {end}\n{text}\n")
    return "\n".join(lines).encode("utf-8")


def export_as_csv(transcriptions: list, include_summary: bool = False) -> bytes:
    """Export a list of transcription records as CSV."""
    output = io.StringIO()
    if not transcriptions:
        return b""
    fields = [
        "id",
        "original_filename",
        "model_used",
        "language",
        "duration_seconds",
        "word_count",
        "status",
        "created_at",
        "transcript",
    ]
    if include_summary:
        fields.append("summary_text")
    writer = csv.DictWriter(output, fieldnames=fields, extrasaction="ignore")
    writer.writeheader()
    for row in transcriptions:
        if include_summary and "summary_text" not in row and "summary" in row:
            row = dict(row)
            row["summary_text"] = row.get("summary", "")
        writer.writerow(row)
    return output.getvalue().encode("utf-8")


def export_as_docx(
    transcript: str,
    title: str = "",
    metadata: dict = None,
    summary_text: str = "",
    include_summary: bool = False,
) -> bytes:
    """Export as DOCX using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document()

    # Title
    heading = doc.add_heading(title or "Transcription", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata table
    if metadata:
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        meta_items = [
            ("Model", metadata.get("model_used", "")),
            ("Language", metadata.get("language", "")),
            ("Duration", f"{metadata.get('duration_seconds', 0):.0f}s"),
            ("Words", str(metadata.get("word_count", 0))),
            ("Exported", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ]
        for key, val in meta_items:
            row = table.add_row()
            row.cells[0].text = key
            row.cells[1].text = val

    doc.add_paragraph()  # spacer

    # Transcript
    doc.add_heading("Transcript", level=2)
    for paragraph in transcript.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())

    if include_summary and summary_text:
        doc.add_paragraph()
        doc.add_heading("Summary", level=2)
        for paragraph in summary_text.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_as_markdown(
    transcript: str,
    title: str = "",
    metadata: dict = None,
    summary_text: str = "",
    include_summary: bool = False,
) -> bytes:
    lines = [f"# {title}\n"] if title else []
    if metadata:
        lines.append("## Metadata\n")
        lines.append(f"- **Model**: {metadata.get('model_used', '')}")
        lines.append(f"- **Language**: {metadata.get('language', '')}")
        dur = metadata.get("duration_seconds", 0)
        lines.append(f"- **Duration**: {int(dur//60)}m {int(dur%60)}s")
        lines.append(f"- **Word Count**: {metadata.get('word_count', 0)}")
        lines.append(f"- **Exported**: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        lines.append("")
    lines.append("## Transcript\n")
    lines.append(transcript)
    if include_summary and summary_text:
        lines.append("")
        lines.append("## Summary\n")
        lines.append(summary_text)
    return "\n".join(lines).encode("utf-8")
